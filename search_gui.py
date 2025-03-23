import os
import re
import time
from pathlib import Path
import fnmatch
import tkinter as tk
from tkinter import filedialog, scrolledtext, ttk, messagebox

# Configuration settings
DIR = "/path/to/your/search/directory"
OUTPUT_DIR = "/path/to/your/output/directory"
DEFAULT_TERMS = ["apple"]
OUTPUT_FILE_TYPES = [".rtf", ".txt", ".md", ".docx"]
DEFAULT_OUTPUT_FILE_TYPE = ".rtf"
EXCERPT_SENTENCES = 5
PROXIMITY_WINDOW = 5
IGNORE_STRING = r"(ignore these patterns)"
IGNORE_FILES = ["index.txt", "*.log"]
IGNORE_FOLDERS = ["temp", "logs"]
MIDDLE_SENTENCES = 10
MIDDLE_WORD_LIMIT = 150
UPDATE_INTERVAL = 50
DOCX_BATCH_SIZE = 1000  # Save .docx every 1000 files
RTF_HEADER = r"{\rtf1\ansi\ansicpg1252\deff0\nouicompat\deflang1033{\fonttbl{\f0\fswiss\fcharset0 Calibri;}}{\colortbl;\red255\green0\blue0;\red0\green0\blue255;}\f0\fs22\par"
RTF_FOOTER = r"}"

# Highlight styles mapping to RTF tags (only used for .rtf)
HIGHLIGHT_STYLES = {
    "Bold": r"\b {term}\b0 ",
    "Red": r"\cf1 {term}\cf0 ",
    "Blue": r"\cf2 {term}\cf0 ",
    "Bold Red": r"\b \cf1 {term}\cf0\b0 ",
    "Bold Blue": r"\b \cf2 {term}\cf0\b0 "
}

# Search modes
SEARCH_MODES = ["Individual Mode", "Proximity Mode"]

# Global state
output_files = {}
start_time = 0
files_processed = 0
total_matches_by_term = {}
total_files = 0
running = False

class Tooltip:
    def __init__(self, widget, text):
        self.widget = widget
        self.text = text
        self.tip_window = None
        self.widget.bind("<Enter>", self.show_tip)
        self.widget.bind("<Leave>", self.hide_tip)

    def show_tip(self, event):
        if self.tip_window or not self.text:
            return
        x = self.widget.winfo_rootx() + 20
        y = self.widget.winfo_rooty() + 20
        self.tip_window = tw = tk.Toplevel(self.widget)
        tw.wm_overrideredirect(True)
        tw.wm_geometry(f"+{x}+{y}")
        label = tk.Label(tw, text=self.text, justify=tk.LEFT,
                         background="#ffffe0", relief=tk.SOLID, borderwidth=1)
        label.pack()

    def hide_tip(self, event):
        if self.tip_window:
            self.tip_window.destroy()
            self.tip_window = None

class SearchApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Text Search Tool")
        self.search_terms = tk.StringVar(value=",".join(DEFAULT_TERMS))
        self.search_dir = tk.StringVar(value=DIR)
        self.output_dir = tk.StringVar(value=OUTPUT_DIR)
        self.highlight_style = tk.StringVar(value="Bold")
        self.output_file_type = tk.StringVar(value=DEFAULT_OUTPUT_FILE_TYPE)
        self.search_mode = tk.StringVar(value=SEARCH_MODES[0])
        self.proximity_window = tk.StringVar(value=str(PROXIMITY_WINDOW))
        self.excerpt_sentences = tk.StringVar(value=str(EXCERPT_SENTENCES))
        self.middle_word_limit = tk.StringVar(value=str(MIDDLE_WORD_LIMIT))
        self.ignore_files = tk.StringVar(value=",".join(IGNORE_FILES))
        self.ignore_folders = tk.StringVar(value=",".join(IGNORE_FOLDERS))
        self.case_sensitive = tk.BooleanVar(value=False)  # Default to case-insensitive
        self.show_middle_excerpt = tk.BooleanVar(value=True)  # Default to showing middle excerpt
        self.create_widgets()

    def create_widgets(self):
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)

        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(5, weight=1)

        # Input Section
        input_frame = ttk.LabelFrame(main_frame, text="Search Inputs", padding="5")
        input_frame.grid(row=0, column=0, padx=5, pady=5, sticky=(tk.W, tk.E))
        tk.Label(input_frame, text="Search Terms (comma-separated):").grid(row=0, column=0, pady=2, sticky=tk.W)
        tk.Entry(input_frame, textvariable=self.search_terms, width=50).grid(row=0, column=1, pady=2)
        tk.Label(input_frame, text="Search Directory:").grid(row=1, column=0, pady=2, sticky=tk.W)
        dir_frame = ttk.Frame(input_frame)
        dir_frame.grid(row=1, column=1, pady=2, sticky=(tk.W, tk.E))
        tk.Entry(dir_frame, textvariable=self.search_dir, width=40).pack(side=tk.LEFT)
        tk.Button(dir_frame, text="Browse", command=self.browse_search_dir).pack(side=tk.LEFT, padx=5)
        tk.Label(input_frame, text="Output Directory:").grid(row=2, column=0, pady=2, sticky=tk.W)
        out_frame = ttk.Frame(input_frame)
        out_frame.grid(row=2, column=1, pady=2, sticky=(tk.W, tk.E))
        tk.Entry(out_frame, textvariable=self.output_dir, width=40).pack(side=tk.LEFT)
        tk.Button(out_frame, text="Browse", command=self.browse_output_dir).pack(side=tk.LEFT, padx=5)
        help_input = tk.Label(input_frame, text="?", fg="blue", cursor="question_arrow")
        help_input.grid(row=0, column=2, padx=5, sticky=tk.W)
        Tooltip(help_input, "Search Terms: Words to find (e.g., apple, pear).\nSearch Directory: Where to look for files.\nOutput Directory: Where results are saved.")

        # Mode and Customization Section
        mode_frame = ttk.LabelFrame(main_frame, text="Search Settings", padding="5")
        mode_frame.grid(row=1, column=0, padx=5, pady=5, sticky=(tk.W, tk.E))
        tk.Label(mode_frame, text="Search Mode:").grid(row=0, column=0, pady=2, sticky=tk.W)
        mode_combo = ttk.Combobox(mode_frame, textvariable=self.search_mode, values=SEARCH_MODES, state="readonly")
        mode_combo.grid(row=0, column=1, pady=2, sticky=tk.W)
        mode_combo.set(SEARCH_MODES[0])
        tk.Label(mode_frame, text="Proximity Window:").grid(row=0, column=2, pady=2, sticky=tk.W)
        proximity_frame = ttk.Frame(mode_frame)
        proximity_frame.grid(row=0, column=3, pady=2, sticky=(tk.W, tk.E))
        self.proximity_entry = tk.Entry(proximity_frame, textvariable=self.proximity_window, width=10)
        self.proximity_entry.pack(side=tk.LEFT)
        self.proximity_entry.config(state=tk.DISABLED)
        tk.Label(mode_frame, text="Case Sensitive:").grid(row=0, column=4, pady=2, sticky=tk.W)
        tk.Checkbutton(mode_frame, variable=self.case_sensitive).grid(row=0, column=5, pady=2, sticky=tk.W)
        tk.Label(mode_frame, text="Excerpt Sentences Limit:").grid(row=1, column=0, pady=2, sticky=tk.W)
        ttk.Entry(mode_frame, textvariable=self.excerpt_sentences, width=10).grid(row=1, column=1, pady=2, sticky=tk.W)
        tk.Label(mode_frame, text="Middle Word Limit:").grid(row=1, column=2, pady=2, sticky=tk.W)
        ttk.Entry(mode_frame, textvariable=self.middle_word_limit, width=10).grid(row=1, column=3, pady=2, sticky=tk.W)
        help_mode = tk.Label(mode_frame, text="?", fg="blue", cursor="question_arrow")
        help_mode.grid(row=0, column=6, padx=5, sticky=tk.W)
        Tooltip(help_mode, "Search Mode: Individual (find each term) or Proximity (terms near each other).\nProximity Window: Sentences around a match to check (default 5).\nCase Sensitive: Match exact case if checked.\nExcerpt Sentences: How many sentences in keyword excerpt.\nMiddle Word Limit: Max words in middle excerpt.")

        # Ignore Settings Section
        ignore_frame = ttk.LabelFrame(main_frame, text="Ignore Settings", padding="5")
        ignore_frame.grid(row=2, column=0, padx=5, pady=5, sticky=(tk.W, tk.E))
        tk.Label(ignore_frame, text="Ignore Files (comma-separated, e.g., index.txt, *.log):").grid(row=0, column=0, pady=2, sticky=tk.W)
        tk.Entry(ignore_frame, textvariable=self.ignore_files, width=50).grid(row=0, column=1, pady=2)
        tk.Label(ignore_frame, text="Ignore Folders (comma-separated, e.g., temp, logs):").grid(row=1, column=0, pady=2, sticky=tk.W)
        tk.Entry(ignore_frame, textvariable=self.ignore_folders, width=50).grid(row=1, column=1, pady=2)
        help_ignore = tk.Label(ignore_frame, text="?", fg="blue", cursor="question_arrow")
        help_ignore.grid(row=0, column=2, padx=5, sticky=tk.W)
        Tooltip(help_ignore, "Ignore Files: File names/patterns to skip (e.g., *.log).\nIgnore Folders: Folder names to exclude (e.g., temp).")

        # Output Settings Section
        output_frame = ttk.LabelFrame(main_frame, text="Output Settings (RTF Output Recommended)", padding="5")
        output_frame.grid(row=3, column=0, padx=5, pady=5, sticky=(tk.W, tk.E))
        tk.Label(output_frame, text="Highlight Style (for RTF only):").grid(row=0, column=0, pady=2, sticky=tk.W)
        style_combo = ttk.Combobox(output_frame, textvariable=self.highlight_style, values=list(HIGHLIGHT_STYLES.keys()), state="readonly")
        style_combo.grid(row=0, column=1, pady=2, sticky=tk.W)
        style_combo.set("Bold")
        tk.Label(output_frame, text="Output File Type:").grid(row=1, column=0, pady=2, sticky=tk.W)
        filetype_combo = ttk.Combobox(output_frame, textvariable=self.output_file_type, values=OUTPUT_FILE_TYPES, state="readonly")
        filetype_combo.grid(row=1, column=1, pady=2, sticky=tk.W)
        filetype_combo.set(DEFAULT_OUTPUT_FILE_TYPE)
        tk.Label(output_frame, text="Show Middle Excerpt:").grid(row=2, column=0, pady=2, sticky=tk.W)
        tk.Checkbutton(output_frame, variable=self.show_middle_excerpt).grid(row=2, column=1, pady=2, sticky=tk.W)
        help_output = tk.Label(output_frame, text="?", fg="blue", cursor="question_arrow")
        help_output.grid(row=0, column=2, padx=5, sticky=tk.W)
        Tooltip(help_output, "Highlight Style: How matches appear in RTF (e.g., Bold, Red)—RTF recommended for speed.\nOutput File Type: .rtf (fast, formatted), .txt/.md (fast, plain), .docx (slower, formatted).\nShow Middle Excerpt: Include middle file context in output if checked.")

        # Button Section
        btn_frame = ttk.Frame(main_frame)
        btn_frame.grid(row=4, column=0, pady=10)
        self.start_btn = tk.Button(btn_frame, text="Start Search", command=self.start_search)
        self.start_btn.pack(side=tk.LEFT, padx=5)
        self.stop_btn = tk.Button(btn_frame, text="Stop", command=self.stop_search, state=tk.DISABLED)
        self.stop_btn.pack(side=tk.LEFT, padx=5)

        # Stats Section
        stats_frame = ttk.LabelFrame(main_frame, text="Search Stats", padding="5")
        stats_frame.grid(row=5, column=0, padx=5, pady=5, sticky=(tk.W, tk.E, tk.N, tk.S))
        stats_frame.columnconfigure(0, weight=1)
        stats_frame.rowconfigure(0, weight=1)
        self.stats_text = scrolledtext.ScrolledText(stats_frame, height=1)
        self.stats_text.grid(row=0, column=0, padx=5, pady=5, sticky=(tk.W, tk.E, tk.N, tk.S))

        self.search_mode.trace("w", self.toggle_proximity_input)
        self.root.update_idletasks()

    def toggle_proximity_input(self, *args):
        if self.search_mode.get() == "Proximity Mode":
            self.proximity_entry.config(state=tk.NORMAL)
        else:
            self.proximity_entry.config(state=tk.DISABLED)

    def browse_search_dir(self):
        dir_path = filedialog.askdirectory(initialdir=self.search_dir.get())
        if dir_path:
            self.search_dir.set(dir_path)

    def browse_output_dir(self):
        dir_path = filedialog.askdirectory(initialdir=self.output_dir.get())
        if dir_path:
            self.output_dir.set(dir_path)

    def update_stats(self, terms):
        self.stats_text.delete(1.0, tk.END)
        self.stats_text.insert(tk.END, f"Searching For: {' '.join(terms)}\n")
        self.stats_text.insert(tk.END, f"Files Processed: {files_processed}/{total_files}\n")
        if self.search_mode.get() == "Individual Mode":
            for term in terms:
                self.stats_text.insert(tk.END, f"Matches for {term}: {total_matches_by_term.get(term, 0)}\n")
        else:
            self.stats_text.insert(tk.END, f"Proximity Matches: {total_matches_by_term.get('proximity', 0)}\n")
        total_time = time.time() - start_time if start_time else 0
        speed = files_processed / total_time if total_time > 0 else 0
        self.stats_text.insert(tk.END, f"Speed (files/sec): {speed:.2f}\n")
        self.stats_text.insert(tk.END, f"Elapsed Time (sec): {total_time:.2f}\n")
        files_left = total_files - files_processed
        time_left = files_left / speed if speed > 0 else float('inf')
        self.stats_text.insert(tk.END, f"Est. Time Left (sec): {time_left:.2f}\n" if time_left != float('inf') else "N/A\n")

        # Dynamically adjust the height based on the number of lines
        content = self.stats_text.get(1.0, tk.END).strip()
        line_count = len(content.splitlines())
        # Set a reasonable maximum height to avoid excessive growth
        max_height = 30
        new_height = min(line_count + 1, max_height)  # Add 1 for padding
        self.stats_text.configure(height=new_height)

        # Force the window to resize based on content
        self.root.update_idletasks()
        self.root.update()

    def start_search(self):
        global running, start_time, files_processed, total_matches_by_term, total_files, output_files
        if running:
            return
        running = True
        self.start_btn.config(state=tk.DISABLED)
        self.stop_btn.config(state=tk.NORMAL)
        files_processed = 0
        start_time = time.time()

        terms = [t.strip() for t in self.search_terms.get().split(',')] if self.search_terms.get() else DEFAULT_TERMS
        flags = 0 if self.case_sensitive.get() else re.IGNORECASE
        self.term_patterns = [re.compile(rf"(?:^|\s){re.escape(term)}(?=[,.\s]|$)", flags) for term in terms]
        total_matches_by_term = {term: 0 for term in terms} if self.search_mode.get() == "Individual Mode" else {'proximity': 0}

        search_dir = self.search_dir.get()
        output_dir = self.output_dir.get()
        output_file_type = self.output_file_type.get()
        if output_file_type != ".rtf" and not hasattr(self, 'warned_non_rtf'):
            self.stats_text.insert(tk.END, "Note: Only .rtf and .docx support highlighting; .docx is slower.\n")
            self.warned_non_rtf = True
        self.is_rtf = output_file_type == ".rtf"
        self.is_docx = output_file_type == ".docx"
        mode = self.search_mode.get()

        try:
            self.excerpt_sentences_val = int(self.excerpt_sentences.get())
            self.middle_word_limit_val = int(self.middle_word_limit.get())
            self.proximity_window_val = int(self.proximity_window.get()) if mode == "Proximity Mode" else PROXIMITY_WINDOW
        except ValueError:
            self.excerpt_sentences_val = EXCERPT_SENTENCES
            self.middle_word_limit_val = MIDDLE_WORD_LIMIT
            self.proximity_window_val = PROXIMITY_WINDOW

        self.ignore_files_list = [f.strip() for f in self.ignore_files.get().split(',')] if self.ignore_files.get() else IGNORE_FILES
        self.ignore_folders_list = [f.strip() for f in self.ignore_folders.get().split(',')] if self.ignore_folders.get() else IGNORE_FOLDERS

        if not os.path.isdir(search_dir):
            self.stats_text.insert(tk.END, f"Error: {search_dir} does not exist\n")
            self.stop_search()
            return

        os.makedirs(output_dir, exist_ok=True)

        overwrite_files = []  # Initialize here
        output_files.clear()  # Reset to avoid bleed-over
        if mode == "Individual Mode":
            for term in terms:
                output_file = os.path.join(output_dir, f"{term}{output_file_type}")
                if os.path.exists(output_file):
                    overwrite_files.append(f"{term}{output_file_type}")
                if self.is_docx:
                    from docx import Document
                    from docx.oxml.ns import qn
                    from docx.oxml import OxmlElement
                    from docx.text.run import Run
                    from docx.shared import RGBColor
                    output_files[term] = Document()
                else:
                    output_files[term] = open(output_file, "w", encoding="utf-8", newline="\n")
                    if self.is_rtf:
                        output_files[term].write(RTF_HEADER)
        elif mode == "Proximity Mode":
            if len(terms) < 2:
                self.stats_text.insert(tk.END, "Error: Proximity Mode requires at least 2 terms\n")
                self.stop_search()
                return
            output_file_name = "_".join(terms).lower() + output_file_type
            output_file = os.path.join(output_dir, output_file_name)
            if os.path.exists(output_file):
                overwrite_files.append(output_file_name)
            if self.is_docx:
                from docx import Document
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                from docx.text.run import Run
                from docx.shared import RGBColor
                output_files['proximity'] = Document()
            else:
                output_files['proximity'] = open(output_file, "w", encoding="utf-8", newline="\n")
                if self.is_rtf:
                    output_files['proximity'].write(RTF_HEADER)

        if overwrite_files:
            warning_msg = "Warning: You're about to overwrite the following existing output files:\n" + "\n".join(overwrite_files) + "\n\nContinue?"
            if not messagebox.askyesno("Overwrite Warning", warning_msg):
                self.stop_search()
                return

        self.txt_files = []
        for pattern in ["*.[tT][xX][tT]", "*.[mM][dD]", "*"]:
            files = list(Path(search_dir).rglob(pattern))
            self.txt_files.extend(f for f in files if f.is_file() and (f.suffix.lower() in ['.txt', '.md'] or not f.suffix))
        self.txt_files = list(set(self.txt_files))
        self.txt_files = [f for f in self.txt_files
                          if not (any(fnmatch.fnmatch(f.name, ignore) for ignore in self.ignore_files_list) or
                                  any(fnmatch.fnmatch(str(f.parent.name), ignore) for ignore in self.ignore_folders_list))]
        total_files = len(self.txt_files)
        self.docx_file_count = 0
        self.update_stats(terms)
        self.root.after(10, self.search_loop)
        
    def search_loop(self):
        global files_processed
        ignore_pattern = re.compile(IGNORE_STRING)
        mode = self.search_mode.get()
        terms = [t.strip() for t in self.search_terms.get().split(',')] if self.search_terms.get() else DEFAULT_TERMS

        for file in self.txt_files:
            if not running:
                break
            file_start = time.time()
            try:
                with open(file, "r", encoding="utf-8", errors="ignore") as f:
                    raw_text = f.read().replace(r'\c', r'\\c')
                    lines = raw_text.splitlines()
                    sentences_all = re.split(r'\.\s*', raw_text)
                    sentences_all = [s.strip() + "." for s in sentences_all if s.strip()]
            except Exception as e:
                self.stats_text.insert(tk.END, f"Error reading {file}: {e}\n")
                files_processed += 1
                continue

            if mode == "Individual Mode":
                for i, line in enumerate(lines, 1):
                    if not running:
                        break
                    if ignore_pattern.search(line):
                        continue
                    for term_idx, pattern in enumerate(self.term_patterns):
                        term = terms[term_idx]
                        match = pattern.search(line)
                        if match:
                            total_matches_by_term[term] += 1
                            start = max(0, i - self.proximity_window_val)
                            end = min(len(lines), i + self.proximity_window_val)
                            excerpt_lines = [l for l in lines[start:end] if not ignore_pattern.search(l)]
                            excerpt_full = " ".join(excerpt_lines)
                            sentences = re.split(r'\.\s*', excerpt_full)
                            keyword_sentence_idx = -1
                            for idx, sentence in enumerate(sentences):
                                if not running:
                                    break
                                if pattern.search(sentence):
                                    keyword_sentence_idx = idx
                                    break
                            if keyword_sentence_idx == -1:
                                keyword_excerpt = line.strip() + "."
                            else:
                                excerpt_start = max(0, keyword_sentence_idx - (self.excerpt_sentences_val // 2))
                                excerpt_end = min(len(sentences), excerpt_start + self.excerpt_sentences_val)
                                keyword_excerpt = " ".join(sentences[excerpt_start:excerpt_end]).strip() + "."
                                if not pattern.search(keyword_excerpt):
                                    keyword_excerpt = line.strip() + "."
                            matched_term = match.group(0)
                            if self.is_rtf:
                                keyword_excerpt = keyword_excerpt.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
                                style = self.highlight_style.get()
                                highlighted_term = HIGHLIGHT_STYLES[style].format(term=matched_term)
                                excerpt_match = pattern.search(keyword_excerpt)
                                if excerpt_match:
                                    start_pos = excerpt_match.start()
                                    end_pos = excerpt_match.end()
                                    keyword_excerpt = (
                                        keyword_excerpt[:start_pos] +
                                        highlighted_term +
                                        keyword_excerpt[end_pos:]
                                    )

                            middle_excerpt = ""
                            if self.show_middle_excerpt.get():
                                all_words = raw_text.split()
                                mid_point = len(all_words) // 2
                                half_limit = self.middle_word_limit_val // 2
                                mid_start = max(0, mid_point - half_limit)
                                mid_end = min(len(all_words), mid_point + half_limit)
                                middle_words = all_words[mid_start:mid_end]
                                middle_excerpt = " ".join(middle_words)
                                if mid_end < len(all_words):
                                    middle_excerpt += "..."
                                if mid_start > 0:
                                    middle_excerpt = "..." + middle_excerpt
                                if self.is_rtf:
                                    middle_excerpt = middle_excerpt.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')

                            out_f = output_files[term]
                            if self.is_docx:
                                from docx.shared import RGBColor
                                keyword_excerpt = ''.join(c for c in keyword_excerpt if ord(c) >= 32 or c in '\t\n\r')
                                middle_excerpt = ''.join(c for c in middle_excerpt if ord(c) >= 32 or c in '\t\n\r')
                                p = out_f.add_paragraph()
                                p.add_run(f"File: {file}", style=None).underline = True
                                p = out_f.add_paragraph()
                                p.add_run("(keyword excerpt):").font.color.rgb = RGBColor(255, 0, 0)
                                p = out_f.add_paragraph()
                                match = pattern.search(keyword_excerpt)
                                if match:
                                    start, end = match.start(), match.end()
                                    pre_text = keyword_excerpt[:start]
                                    matched_term = keyword_excerpt[start:end]
                                    post_text = keyword_excerpt[end:]
                                    if pre_text:
                                        p.add_run(pre_text)
                                    run = p.add_run(matched_term)
                                    style = self.highlight_style.get()
                                    if "Bold" in style:
                                        run.bold = True
                                    if "Red" in style:
                                        run.font.color.rgb = RGBColor(255, 0, 0)
                                    elif "Blue" in style:
                                        run.font.color.rgb = RGBColor(0, 0, 255)
                                    if post_text:
                                        p.add_run(post_text)
                                else:
                                    p.add_run(keyword_excerpt)
                                if self.show_middle_excerpt.get():
                                    p = out_f.add_paragraph()
                                    p.add_run("Middle of file excerpt:").font.color.rgb = RGBColor(255, 0, 0)
                                    out_f.add_paragraph(middle_excerpt)
                                out_f.add_paragraph("------------------------")
                                # Batch save on matches
                                if total_matches_by_term[term] % DOCX_BATCH_SIZE == 0:
                                    from docx import Document
                                    temp_file = os.path.join(self.output_dir.get(), f"temp_{term}_{total_matches_by_term[term] // DOCX_BATCH_SIZE}.docx")
                                    out_f.save(temp_file)
                                    output_files[term] = Document()
                            elif self.is_rtf:
                                out_f.write(f"\\ul File: {file}\\ulnone\\par\n")
                                out_f.write("\\par\n")
                                out_f.write(f"\\cf1 (keyword excerpt):\\cf0\\par\n")
                                out_f.write(f"{keyword_excerpt}\\par\n")
                                if self.show_middle_excerpt.get():
                                    out_f.write("\\par\n")
                                    out_f.write(f"\\cf1 Middle of file excerpt:\\cf0\\par\n")
                                    out_f.write(f"{middle_excerpt}\\par\n")
                                out_f.write("------------------------\\par\n")
                                out_f.write("\\par\n")
                            else:
                                out_f.write(f"File: {file}\n")
                                out_f.write("\n")
                                out_f.write(f"(keyword excerpt):\n")
                                out_f.write(f"{keyword_excerpt}\n")
                                if self.show_middle_excerpt.get():
                                    out_f.write("\n")
                                    out_f.write(f"Middle of file excerpt:\n")
                                    out_f.write(f"{middle_excerpt}\n")
                                out_f.write("------------------------\n")
                            if not self.is_docx:
                                out_f.flush()
            else:
                proximity_matches = []
                for i, sentence in enumerate(sentences_all):
                    if not running:
                        break
                    if ignore_pattern.search(sentence):
                        continue
                    if self.term_patterns[0].search(sentence):
                        start = max(0, i - self.proximity_window_val)
                        end = min(len(sentences_all), i + self.proximity_window_val + 1)
                        window = " ".join(sentences_all[start:end])
                        if all(pattern.search(window) for pattern in self.term_patterns):
                            proximity_matches.append(i)

                for match_idx in proximity_matches:
                    if not running:
                        break
                    total_matches_by_term['proximity'] += 1
                    start = max(0, match_idx - (self.excerpt_sentences_val // 2))
                    end = min(len(sentences_all), start + self.excerpt_sentences_val)
                    excerpt_sentences_list = sentences_all[start:end]
                    keyword_excerpt = " ".join(excerpt_sentences_list)

                    if not all(pattern.search(keyword_excerpt) for pattern in self.term_patterns):
                        continue

                    if self.is_rtf:
                        keyword_excerpt = keyword_excerpt.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
                        style = self.highlight_style.get()
                        for pattern_idx, pattern in enumerate(self.term_patterns):
                            matches = list(pattern.finditer(keyword_excerpt))
                            for match in reversed(matches):
                                if not running:
                                    break
                                matched_term = match.group(0)
                                highlighted_term = HIGHLIGHT_STYLES[style].format(term=matched_term)
                                keyword_excerpt = (
                                    keyword_excerpt[:match.start()] +
                                    highlighted_term +
                                    keyword_excerpt[match.end():]
                                )

                    middle_excerpt = ""
                    if self.show_middle_excerpt.get():
                        all_words = raw_text.split()
                        mid_point = len(all_words) // 2
                        half_limit = self.middle_word_limit_val // 2
                        mid_start = max(0, mid_point - half_limit)
                        mid_end = min(len(all_words), mid_point + half_limit)
                        middle_words = all_words[mid_start:mid_end]
                        middle_excerpt = " ".join(middle_words)
                        if mid_end < len(all_words):
                            middle_excerpt += "..."
                        if mid_start > 0:
                            middle_excerpt = "..." + middle_excerpt
                        if self.is_rtf:
                            middle_excerpt = middle_excerpt.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')

                    out_f = output_files['proximity']
                    if self.is_docx:
                        from docx.shared import RGBColor
                        keyword_excerpt = ''.join(c for c in keyword_excerpt if ord(c) >= 32 or c in '\t\n\r')
                        middle_excerpt = ''.join(c for c in middle_excerpt if ord(c) >= 32 or c in '\t\n\r')
                        p = out_f.add_paragraph()
                        p.add_run(f"File: {file}", style=None).underline = True
                        p = out_f.add_paragraph()
                        p.add_run("(keyword excerpt):").font.color.rgb = RGBColor(255, 0, 0)
                        p = out_f.add_paragraph()
                        current_text = keyword_excerpt
                        style = self.highlight_style.get()
                        for pattern in self.term_patterns:
                            match = pattern.search(current_text)
                            if match:
                                start, end = match.start(), match.end()
                                pre_text = current_text[:start]
                                matched_term = current_text[start:end]
                                post_text = current_text[end:]
                                if pre_text:
                                    p.add_run(pre_text)
                                run = p.add_run(matched_term)
                                if "Bold" in style:
                                    run.bold = True
                                if "Red" in style:
                                    run.font.color.rgb = RGBColor(255, 0, 0)
                                elif "Blue" in style:
                                    run.font.color.rgb = RGBColor(0, 0, 255)
                                current_text = post_text
                            else:
                                break
                        if current_text:
                            p.add_run(current_text)
                        if self.show_middle_excerpt.get():
                            p = out_f.add_paragraph()
                            p.add_run("Middle of file excerpt:").font.color.rgb = RGBColor(255, 0, 0)
                            out_f.add_paragraph(middle_excerpt)
                        out_f.add_paragraph("------------------------")
                        # Batch save on matches
                        if total_matches_by_term['proximity'] % DOCX_BATCH_SIZE == 0:
                            from docx import Document
                            temp_file = os.path.join(self.output_dir.get(), f"temp_proximity_{total_matches_by_term['proximity'] // DOCX_BATCH_SIZE}.docx")
                            out_f.save(temp_file)
                            output_files['proximity'] = Document()
                    elif self.is_rtf:
                        out_f.write(f"\\ul File: {file}\\ulnone\\par\n")
                        out_f.write("\\par\n")
                        out_f.write(f"\\cf1 (keyword excerpt):\\cf0\\par\n")
                        out_f.write(f"{keyword_excerpt}\\par\n")
                        if self.show_middle_excerpt.get():
                            out_f.write("\\par\n")
                            out_f.write(f"\\cf1 Middle of file excerpt:\\cf0\\par\n")
                            out_f.write(f"{middle_excerpt}\\par\n")
                        out_f.write("------------------------\\par\n")
                        out_f.write("\\par\n")
                    else:
                        out_f.write(f"File: {file}\n")
                        out_f.write("\n")
                        out_f.write(f"(keyword excerpt):\n")
                        out_f.write(f"{keyword_excerpt}\n")
                        if self.show_middle_excerpt.get():
                            out_f.write("\n")
                            out_f.write(f"Middle of file excerpt:\n")
                            out_f.write(f"{middle_excerpt}\n")
                        out_f.write("------------------------\n")
                    if not self.is_docx:
                        out_f.flush()

            files_processed += 1
            if files_processed % UPDATE_INTERVAL == 0:
                elapsed = time.time() - file_start
                self.stats_text.insert(tk.END, f"Processed {file} in {elapsed:.2f} seconds\n")
                self.update_stats(terms)
                self.root.update()
                if not running:
                    break

        self.finalize_search(terms)

    def stop_search(self):
        global running
        if running:
            running = False
            self.stats_text.insert(tk.END, "Stopping search...\n")
        self.start_btn.config(state=tk.NORMAL)
        self.stop_btn.config(state=tk.DISABLED)

    def finalize_search(self, terms):
        global output_files, running
        mode = self.search_mode.get()
        if self.is_docx:
            from docx import Document
            # Save and merge all in-memory batches
            for key, f in list(output_files.items()):
                if f:
                    total_matches = total_matches_by_term.get(key, 0)
                    batch_count = (total_matches // DOCX_BATCH_SIZE) + (1 if total_matches % DOCX_BATCH_SIZE else 0)
                    temp_file = os.path.join(self.output_dir.get(), f"temp_{key}_{batch_count}.docx")
                    f.save(temp_file)
                    # Merge all temp files for this key
                    final_file = os.path.join(self.output_dir.get(), f"{key}.docx" if key != 'proximity' else "_".join(terms).lower() + ".docx")
                    merged_doc = Document()
                    for i in range(1, batch_count + 1):
                        batch_file = os.path.join(self.output_dir.get(), f"temp_{key}_{i}.docx")
                        if os.path.exists(batch_file):
                            temp_doc = Document(batch_file)
                            for element in temp_doc.element.body:
                                merged_doc.element.body.append(element)
                    merged_doc.save(final_file)
            # Clean up all temp files
            keys_to_clean = terms if mode == "Individual Mode" else ['proximity']
            for key in keys_to_clean:
                total_matches = total_matches_by_term.get(key, 0)
                batch_count = (total_matches // DOCX_BATCH_SIZE) + (1 if total_matches % DOCX_BATCH_SIZE else 0)
                for i in range(1, batch_count + 1):
                    temp_file = os.path.join(self.output_dir.get(), f"temp_{key}_{i}.docx")
                    if os.path.exists(temp_file):
                        os.remove(temp_file)
        else:
            for key, f in output_files.items():
                if f and not f.closed:
                    if self.is_rtf:
                        f.write(RTF_FOOTER)
                    f.close()
        self.update_stats(terms)
        if running:
            self.stats_text.insert(tk.END, "\nSearch Completed Successfully\n")
            running = False
        else:
            self.stats_text.insert(tk.END, "\nSearch Stopped by User\n")
        self.start_btn.config(state=tk.NORMAL)
        self.stop_btn.config(state=tk.DISABLED)
        
if __name__ == "__main__":
    root = tk.Tk()
    app = SearchApp(root)
    try:
        root.mainloop()
    except KeyboardInterrupt:
        app.stop_search()
        app.stats_text.insert(tk.END, "Search interrupted by user\n")
        for key, f in output_files.items():
            if f:
                if app.is_rtf:
                    if not f.closed:
                        f.write(RTF_FOOTER)
                        f.close()
                elif app.is_docx:
                    f.save(os.path.join(app.output_dir.get(), f"{key}.docx" if key != 'proximity' else "_".join(app.search_terms.get().split(',')).lower() + ".docx"))
                else:
                    if not f.closed:
                        f.close()
        root.quit()
    except ImportError as e:
        if "docx" in str(e):
            app.stats_text.insert(tk.END, "Error: .docx support requires 'pip install python-docx'\n")
        raise
